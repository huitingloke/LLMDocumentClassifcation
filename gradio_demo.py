import gradio as gr
import os
from pypdf import PdfReader
import chromadb
import uuid
import datetime
from ollama import ChatResponse, chat
from docx import Document
import csv
import io
import json
import time
<<<<<<< Updated upstream
=======
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.llms import OpenAI 
from dotenv import dotenv_values
import langchain_implementation

>>>>>>> Stashed changes

length_of_text = 6000
# Category Tag choices
choices = [
    "Regulation",
    "Notices, News",
    "Financial Data/Reports",
    "Client Info",
    "Constitution",
    "Contracts",
    "T&Cs",
    "Privacy Policy",
    "Own Financial Data & Reports",
]

client = chromadb.PersistentClient(path="./chromadb_persistent_storage/")

# FOR TESTING PURPOSES
client.delete_collection("document_storage")
collection = client.get_or_create_collection(
    name="document_storage", 
)

def navigate_home():
    """Returns to the home page by toggling visibility."""
    return gr.update(visible=True), gr.update(visible=False)

def navigate_classification():
    """Returns to the classification page by toggling visibility."""
    return gr.update(visible=True), gr.update(visible=False)

def navigate_search():
    """Returns to the search page by toggling visibility."""
    return gr.update(visible=True), gr.update(visible=False)

def navigate_savedDocuments():
    """Returns to the search page by toggling visibility."""
    return gr.update(visible=True), gr.update(visible=False)

def show_page(page):
    return {
        "home": [gr.update(visible=True), gr.update(visible=False), gr.update(visible=False), gr.update(visible=False)],
        "classification": [gr.update(visible=False), gr.update(visible=True), gr.update(visible=False), gr.update(visible=False)],
        "search": [gr.update(visible=False), gr.update(visible=False), gr.update(visible=True), gr.update(visible=False)],
        "savedDocuments": [gr.update(visible=False), gr.update(visible=False), gr.update(visible=False), gr.update(visible=True)],
    }[page]

def upload_file(file, notes): #Another function exists below, KIV
    if file is not None:
        file_path = file.name  # Save file path for reuse
        file_name = os.path.basename(file.name)  # Extract just the file name
        status = f"✅ File '{file_name}' uploaded successfully!\n\n📌 Notes: {notes if notes else 'No additional notes'}"
        return status, file_path, gr.update(visible=False), gr.update(visible=True)
    else:
        return "❌ No file uploaded.", None

## For now, document preview only supports PDF files ##
def preview_document(file):
    """Extracts text preview from PDF files (first page only)."""
    if file is None:
        return "No file uploaded."
    
    if file.name.endswith(".pdf"):
        with open(file.name, "rb") as f:  # Open the file explicitly
            reader = PdfReader(f)
            text = reader.pages[0].extract_text() if reader.pages else "Empty PDF"
        return text[:500]  # Show only first 500 characters
    
    return "Unsupported file type."

def search_by_tag_and_query(query, filters_contentType, filters_authors, filters_postedAt):
    results = collection.query(
        query_texts=[query],
        n_results=10000000000000000000,
        where={"level2": filters_contentType},
    )
    return results

"""def retrieve_search(query, num_results):

    print(f"Retrieving search results...: QUERY: {query}; NUM_RESULTS: {num_results}")
    results = ""
    num_results = int(num_results)
    collection_retrieval = collection.query(
        query_texts=[query],
        n_results=num_results,
    )

    if collection_retrieval:
        for item in collection_retrieval:
            print(item)
    return collection_retrieval if collection_retrieval else "No results found.""" 

def extract_text(file):

    print(file, file.name)
    if file.name.lower().endswith(".pdf"):
        print("Filetype: PDF", file.name)
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        print(text)
        return text
    
    elif file.name.lower().endswith(".txt") or file.name.endswith(".csv"):
        print("Filetype: txt")
        return file.read().decode("utf-8")
    
    elif file.name.lower().endswith(".docx") or file.name.endswith(".doc"):
        print("Filetype: Word", file.name)
        document = Document(file)
        text = ""
        for para in document.paragraphs:
            text += para.text
        print(type(text))
        return text
        
    elif file.name.lower().endswith(".html"): 
        print("Filetype: HTML", file.name)
        return "WIP"
        #return file.read().decode("utf-8")

    elif file.name.lower().endswith(".json"):
        the_file = file.read().decode("utf-8")
        json_data = json.loads(the_file)
        output = io.StringIO()
        
        if json_data:
            writer = csv.DictWriter(output, fieldnames=json_data[0].keys())
            writer.writeheader()
            writer.writerows(json_data)
    
        return output.getvalue()
    
    else: 
        return "Unsupported file type"

def generate_response(document_text, chosen_model):

    prompt = """
    Given the following document text, classify it into one of the following categories:

    Level 1 Categories: 
    - internal
    - external

    If the document falls under internal, further classify it into:
    Level 2 Categories (Internal): 
    - Constitution
    - Contracts
    - T&Cs
    - Privacy Policy
    - Own Financial Data & Reports

    If the document falls under external, further classify it into:
    Level 2 Categories (External): 
    - Regulation
    - Notices, News
    - Financial Data/Reports
    - Client Info

    Only return the classification in the following string format:
    ```json
    {
      "level_1_category": "internal" OR "external",
      "level_2_category": "[topic]"
    }
    ```

    Do not elaborate on your response. It must only be JSON. 

    **Document Text**:
    """

    response: ChatResponse = chat(model=chosen_model, messages=[
        {
            'role': 'user',
            'content': f"{prompt}: {document_text}",
        },
    ])
    # or access fields directly from the response object

    model_response = response.message.content
    try:

        filtered_response = model_response.split("```")[1] if "```" in model_response else model_response
        filtered_response = model_response.split("json")[1].lstrip().split("```")[0].strip() if "json" in model_response else model_response 
        
    except: 
        filtered_response = model_response
    print(filtered_response)
    return filtered_response

def process_file(file_uploader, notes="", chosen_model="deepseek-r1:7b"):

    start_time = time.time()
    if file_uploader is not None:
        metadata = {}
        text = extract_text(file_uploader)

        if text == "WIP" or text == "Unsupported file type":
            return "Work in Progress"

        metadata["date"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        metadata["source"] = ""
        metadata["comments"] = notes or "NA"

        classified = generate_response(document_text=text[:length_of_text], chosen_model=chosen_model)
        print("CLASSIFIED: " + classified)
        classified = json.loads(classified)
        metadata["level1"] = classified["level_1_category"]
        metadata["level2"] = classified["level_2_category"]
        metadata["model"] = chosen_model
 
        collection.add(
            documents=[text],
            metadatas=[metadata],
            ids=[uploaded_id:=str(uuid.uuid4())],
        )

        print(metadata, uploaded_id)

        end_time = time.time()
        execution_time = end_time - start_time
        print(f"Execution time: {execution_time} seconds")
        
        return f"File '{file_uploader.name}' uploaded successfully with these comments: {notes}"
    
    return "No file uploaded."

with gr.Blocks(css="""
    @import url('https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0-beta3/css/all.min.css');
    @import url('https://fonts.googleapis.com/css2?family=Manrope:wght@300;400;600;700&display=swap');
    @import url('https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.2/css/all.min.css');

    * {
        font-family: 'Manrope', sans-serif !important;
        font-size: 16px !important;
    }

    /* Logo at top left hand corner */
    .logo {
        font-family: 'Arial', sans-serif !important;  /* Specify font */
        color: white !important;  /* Set text color to white */
        font-size: 40px !important;  /* Set logo font size */
        font-weight: bold !important;  /* Ensure bold font */
    }

    .top-bar {
        display: flex;
        align-items: center;  /* Align items at the center */
        justify-content: space-between;
        background-color: #cc0000;
        padding: 5px 20px;
        border-radius: 0px;
        height: 85px;
    }
    
/* Style for the "Upload Document(s)" button */

    /* Search box */
    .search-box {
        display: flex;
        align-items: center;
        position: relative;
        z-index: 10;
    }

    .search-box input {
        width: 596px;
        padding: 8px;
        border-radius: 25px;
        border: 1px solid #ddd;
    }

    .search-box i {
        position: absolute;
        left: 10px;
        color: #aaa;
        font-size: 20px;
    }

    .gr-button {
        background-color: #cc0000 !important;
        color: white !important;
        border: none !important;
        font-size: 16px !important;
    }

    .gr-button:hover {
        background-color: #990000 !important;
    }

    /* Move the "Log out" button to the bottom left */
    .logout-button-container {
        display: flex;
        justify-content: flex-start;
        position: absolute;
        bottom: 20px;
        left: 20px;
        width: 10%;
    }

    /* Profile section in the top-right corner */
    .profile-section {
        display: flex;
        flex-direction: column;
        color: white;
        align-items: flex-end;
    }

    .profile-section i {
        font-size: 24px;
        margin-bottom: 5px;
    }

    .profile-section .profile-name {
        font-size: 16px;
        font-weight: bold;
        text-align: right;
    }
    
    .sidebar {
        background-color: #f8f8f8;
        padding: 20px;
        height: 100vh; /* Full height */
        width: 200px;
        font-weight: bold;
    }

    .sidebar-item {
        padding: 10px;
        font-size: 16px;
        color: black;
        cursor: pointer; /* Makes it look clickable */
    }

    .sidebar-item:hover {
        color: blue; /* Changes color on hover */
    }

    .search-button-container {
        width: 160px;
    }      

    .search-box {
        width: 800px;
    }   
    

""") as demo:

    # Top Bar
    with gr.Row(elem_classes="top-bar"):
        gr.Markdown("**NOMURA**", elem_classes="logo")

        # Profile section in the top-right corner
        gr.HTML("""
            <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.2/css/all.min.css">
            <div class="profile-section">             
                <span class="profile-name">Welcome back!<br>Xavier Lee</span>
                <i class="fas fa-user-circle"></i>  <!-- Profile icon -->
            </div>
        """)

    # Define home and classification pages
    home_page = gr.Column(visible=True)
    classification_page = gr.Column(visible=False)
    search_page = gr.Column(visible=False)
    savedDocuments_page = gr.Column(visible=False)
    existing_page = gr.Column(visible=False)

    # HOME 
    with home_page:
        with gr.Row():
            with gr.Column(scale=1):
                gr.HTML("<div class='nav-item' onclick='document.getElementById('home_btn').click()'>🏠 Home</div>")
                gr.HTML("<div class='nav-item' onclick='document.getElementById('search_btn').click()'>🔍 Search</div>")
                gr.HTML("<div class='sidebar-item' onclick='document.getElementById('savedDocs_btn').click()'>📂 Saved Documents</div>")
                gr.HTML("<div class='sidebar-item' onclick='document.getElementById('settings_btn').click()'>⚙️ Settings</div>")

            with gr.Column(scale=4):
                gr.Markdown("### Upload your Document(s) here")
                file_uploader = gr.File(label="Upload your document", type="filepath")
                chosen_model = gr.Dropdown(choices=["deepseek-r1:7b", "llama3.1", "phi4:14b", "gpt3.5-turbo"], label="Choose a model", value="gpt3.5-turbo", interactive=True)
                notes = gr.Textbox(placeholder="Write any comments you have about your document(s) here.", label="Comments")       
                upload_button = gr.Button("Upload Document(s)")
                output_text = gr.Textbox(label="Upload Status", interactive=False)
                classify_button = gr.Button("Classify Document(s)")

                # Link button to file upload
                upload_button.click(
                    fn=upload_file,
                    inputs=[file_uploader, notes],
                    outputs=[output_text, file_uploader] 
                )

                classify_button.click(navigate_classification, outputs=[classification_page, home_page])
                

    # CLASSIFICATION 
    with classification_page:
        ## Storing does not seem to work ##
        classification_file = gr.State()  # Store file path in state

        with gr.Row():
            with gr.Column(scale=1):
                gr.HTML("<div class='nav-item' onclick='document.getElementById('home_btn').click()'>🏠 Home</div>")
                gr.HTML("<div class='nav-item' onclick='document.getElementById('search_btn').click()'>🔍 Search</div>")
                gr.HTML("<div class='sidebar-item' onclick='document.getElementById('savedDocs_btn').click()'>📂 Saved Documents</div>")
                gr.HTML("<div class='sidebar-item' onclick='document.getElementById('settings_btn').click()'>⚙️ Settings</div>")

            with gr.Column(scale=3):
                gr.Markdown("### Document Preview")
                file_uploader = gr.File(label="Upload Document") # uploaded file from home page should show here
                document_preview = gr.Textbox(label="Document Preview", interactive=False)   
                file_uploader.change(preview_document, inputs=classification_file, outputs=document_preview)           

                with gr.Row():
                    reclassify_button = gr.Button("Reclassify", elem_classes="reclassify-button")
                    backtoHome_button = gr.Button("New classification", elem_classes="backtoHome-button") # reset all fields from prev page automatically

            with gr.Column(scale=1):
                gr.Markdown("### Classification Results")
                classification_contentType = gr.TextArea(label="Content type", interactive=False)
                classification_contentType = gr.TextArea(label="Author(s)", interactive=False)
                classification_contentType = gr.TextArea(label="Posted at", interactive=False)

            # Button to return to home page
            backtoHome_button.click(navigate_home, outputs=[home_page, classification_page])

            # Button to search page to test
            searchPage_button = gr.Button("Search", elem_classes="")
            searchPage_button.click(navigate_search, outputs=[search_page, classification_page])
        

    # SEARCH 
    with search_page:
        with gr.Row():
            with gr.Column(scale=1):
                gr.HTML("<div class='nav-item' onclick='document.getElementById('home_btn').click()'>🏠 Home</div>")
                gr.HTML("<div class='nav-item' onclick='document.getElementById('search_btn').click()'>🔍 Search</div>")
                gr.HTML("<div class='sidebar-item' onclick='document.getElementById('savedDocs_btn').click()'>📂 Saved Documents</div>")
                gr.HTML("<div class='sidebar-item' onclick='document.getElementById('settings_btn').click()'>⚙️ Settings</div>")

            with gr.Column(scale=5):
                with gr.Row():
                    with gr.Column(): 
                        search_input = gr.Textbox(placeholder="Enter your search query", elem_classes="search-box")
                    
                    with gr.Column(): 
                        # filter by content type
                        filters_contentType = gr.Dropdown(choices=["Content Type(s)", "Option 1", "Option 2", "Option 3"], 
                      multiselect=True, elem_classes="filter-box", value="Content Type(s)")
                    with gr.Column(): 
                        # filter by author
                        filters_authors = gr.Dropdown(choices=["Author(s)", "Option 1", "Option 2", "Option 3"], 
                      multiselect=True, elem_classes="filter-box", value="Author(s)")
                    with gr.Column(): 
                        # filter by date
                        filters_postedAt = gr.Dropdown(choices=["Date(s)", "Option 1", "Option 2", "Option 3"], 
                      multiselect=True, elem_classes="filter-box", value="Date(s)")
                    
                    with gr.Column():
                        search_button = gr.Button("Search", elem_classes="search-button-container")

                search_documents_output = gr.TextArea(label="Search Results", interactive=False)


            ## TEMP hardcoded outputs ##
            search_button.click(lambda query, num=None: "Iterate through fake database.", 
                    inputs=[search_input, filters_contentType, filters_authors, filters_postedAt], 
                    outputs=search_documents_output)




    # SAVED DOCUMENTS
    with savedDocuments_page:
        with gr.Row():
            with gr.Column(scale=1):
                gr.HTML("<div class='nav-item' onclick='document.getElementById('home_btn').click()'>🏠 Home</div>")
                gr.HTML("<div class='nav-item' onclick='document.getElementById('search_btn').click()'>🔍 Search</div>")
                gr.HTML("<div class='sidebar-item' onclick='document.getElementById('savedDocs_btn').click()'>📂 Saved Documents</div>")
                gr.HTML("<div class='sidebar-item' onclick='document.getElementById('settings_btn').click()'>⚙️ Settings</div>")
            
            gr.Markdown("### Saved Documents")

    with existing_page:
        with gr.Row():
            with gr.Column(scale=1):
                gr.Markdown("### Home")
                gr.Markdown("### Search")
                gr.Markdown("### Saved Documents")
                gr.Markdown("### Settings")

            with gr.Column(scale=4):
                gr.Markdown("### Upload your Document(s) here")
                file_uploader = gr.File(label="Upload your document", type="filepath")
                chosen_model = gr.Dropdown(choices=["deepseek-r1:7b", "llama3.1", "phi4:14b"], label="Choose a model", value="deepseek-r1:7b")
                notes = gr.Textbox(placeholder="Write any comments you have about your document(s) here.", label="Comments")
                upload_button = gr.Button("Upload Document(s)")
                reset_button = gr.Button("Reset")

                output_text = gr.Textbox(label="Upload Status", interactive=False)

                gr.Markdown("---")
                gr.Markdown("### Search")
                query = gr.Textbox(label="Search", placeholder="Enter your search query")
                category_tag = gr.Dropdown(choices=choices, label="Category Filter", value=choices[0])
                search_button = gr.Button("Search")
                search_documents_output = gr.TextArea(label="Search Results", interactive=False)
                search_button.click(search_by_tag_and_query, inputs=[query, category_tag], outputs=search_documents_output)
                upload_button.click(process_file, inputs=[file_uploader, notes, chosen_model], outputs=output_text)

    # Log out button container at the bottom-left of the page
    with gr.Row(elem_classes="logout-button-container"):
        gr.Button("Log out")

    # Invisible helper buttons for routing
    home_btn = gr.Button(visible=False, elem_id="home_btn")
    search_btn = gr.Button(visible=False, elem_id="search_btn")
    savedDocs_btn = gr.Button(visible=False, elem_id="savedDocs_btn")
    settings_btn = gr.Button(visible=False, elem_id="settings_btn")


    home_btn.click(fn=show_page, inputs=[gr.State("home")], outputs=[home_page, classification_page, search_page, savedDocuments_page])
    search_btn.click(fn=show_page, inputs=[gr.State("search")], outputs=[home_page, classification_page, search_page, savedDocuments_page])
    savedDocs_btn.click(fn=show_page, inputs=[gr.State("savedDocuments")], outputs=[home_page, classification_page, search_page, savedDocuments_page])


# Inject JavaScript to trigger invisible buttons for navigation
with demo:
    gr.HTML("""
    <script>
    document.addEventListener("DOMContentLoaded", function () {
        document.getElementById("home_btn").style.display = "none";
        document.getElementById("search_btn").style.display = "none";
        document.getElementById("savedDocs_btn").style.display = "none";
        document.getElementById("settings_btn").style.display = "none";
    });
    </script>
    """)


if __name__ == "__main__":
    demo.launch()



    # home_btn = gr.Button(visible=False)
    # # classification_btn = gr.Button(visible=False)
    # upload_btn = gr.Button(visible=False)
    # savedDocs_btn = gr.Button(visible=False)

    # <script>
    #     function navigateTo(page) {
    #         if (page === 'home') { document.querySelector('button:nth-of-type(1)').click(); }
    #         if (page === 'search') { document.querySelector('button:nth-of-type(2)').click(); }
    #         if (page === 'savedDocuments') { document.querySelector('button:nth-of-type(3)').click(); }
    #     }
    # </script>

    
