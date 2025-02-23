import gradio as gr
import os
from pypdf import PdfReader
# from docx import Document
import pandas as pd
import database_implementation
import assistive_functions
import langchain_implementation
from docx import Document

choices = [
    "Regulation",
    "Notices, News",
    "Financial Data/Reports",
    "Client Info",
    "Constitution",
    "Contracts",
    "T&Cs",
    "Privacy Policy",
    "Own Financial Data & Reports",
]
def reset_home():
    return None, "", "", ""

def navigate_home():
    """Returns to the home page by toggling visibility."""
    return gr.update(visible=True), gr.update(visible=False), gr.update(visible=False), gr.update(visible=False)

def navigate_classification():
    """Returns to the classification page by toggling visibility."""
    return gr.update(visible=False), gr.update(visible=True), gr.update(visible=False), gr.update(visible=False)

def navigate_search():
    """Returns to the search page by toggling visibility."""
    return gr.update(visible=False), gr.update(visible=False), gr.update(visible=True), gr.update(visible=False)

def navigate_savedDocuments():
    """Returns to the search page by toggling visibility."""
    return gr.update(visible=False), gr.update(visible=False), gr.update(visible=False), gr.update(visible=True)

def upload_file(file, notes):
    if file is not None:
        file_name = os.path.basename(file.name)  # Extract just the file name
        process = assistive_functions.process_file(
            file_name, 
            notes=notes,
        )
        status = f"✅ File '{file_name}' uploaded successfully!\n\n📌 Notes: {notes if notes else 'No additional notes'}" if process else "❌ No file uploaded."
        return status

# def upload_to_classify(file):
#     if file is not None:
#         file_path = file.name  # Save file path for reuse
#         return file_path, gr.update(visible=False), gr.update(visible=True)
#     else:
#         return "❌ No file uploaded.", None, None

def upload_to_classify_preview_document(file):
    """Extracts text preview from PDF, DOCX, CSV, and XLSX files."""
    if file is None:
        return "No file uploaded."
    
    file_name = file.name.lower()

    # PDF Preview
    if file_name.endswith(".pdf"):
        try:
            with open(file.name, "rb") as f:
                reader = PdfReader(f)
                text = reader.pages[0].extract_text() if reader.pages else "Empty PDF"
            return text[:500] , gr.update(visible=False), gr.update(visible=True) # Show only first 500 characters
        except Exception as e:
            return f"Error reading PDF: {e}", gr.update(visible=False), gr.update(visible=True)

    elif file.name.lower().endswith(".docx") or file.name.endswith(".doc"):
        print("Filetype: Word", file.name)
        document = Document(file)
        text = ""
        for para in document.paragraphs:
            text += para.text
        return text[:500]

    # CSV Preview
    elif file_name.endswith(".csv"):
        try:
            df = pd.read_csv(file.name, nrows=5)  # Read first 5 rows
            return df.to_string(index=False), gr.update(visible=False), gr.update(visible=True)  # Convert to string
        except Exception as e:
            return f"Error reading CSV: {e}", gr.update(visible=False), gr.update(visible=True)

    # Excel (XLSX) Preview
    elif file_name.endswith(".xlsx"):
        try:
            df = pd.read_excel(file.name, engine="openpyxl", nrows=5)
            return df.to_string(index=False), gr.update(visible=False), gr.update(visible=True)
        except Exception as e:
            return f"Error reading XLSX: {e}", gr.update(visible=False), gr.update(visible=True)

    return "Unsupported file type."


with gr.Blocks(css="""
    @import url('https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0-beta3/css/all.min.css');
    @import url('https://fonts.googleapis.com/css2?family=Manrope:wght@300;400;600;700&display=swap');
    @import url('https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.2/css/all.min.css');

    * {
        font-family: 'Manrope', sans-serif !important;
        font-size: 16px !important;
    }

    /* Logo at top left hand corner */
    .logo {
        font-family: 'Arial', sans-serif !important;  /* Specify font */
        color: white !important;  /* Set text color to white */
        font-size: 40px !important;  /* Set logo font size */
        font-weight: bold !important;  /* Ensure bold font */
    }

    .top-bar {
        display: flex;
        align-items: center;  /* Align items at the center */
        justify-content: space-between;
        background-color: #cc0000;
        padding: 5px 20px;
        border-radius: 0px;
        height: 85px;
    }

    /* Style for the "Upload Document(s)" button */
    .gr-button {
        background-color: #cc0000 !important;
        color: white !important;
        border: none !important;
        font-size: 16px !important;
    }

    .gr-button:hover {
        background-color: #990000 !important;
    }

    /* Move the "Log out" button to the bottom left */   
    .logout-button-container {
        position: fixed;
        bottom: 20px;
        left: 20px;
        width: 160px;
    }

    /* Profile section in the top-right corner */
    .profile-section {
        display: flex;
        flex-direction: column;
        color: white;
        align-items: flex-end;
    }

    .profile-section i {
        font-size: 24px;
        margin-bottom: 5px;
    }

    .profile-section .profile-name {
        font-size: 16px;
        font-weight: bold;
        text-align: right;
    }
    
    .sidebar {
        background-color: #f8f8f8;
        padding: 20px;
        height: 100vh; /* Full height */
        width: 200px;
        font-weight: bold;
    }

    .sidebar-item {
        padding: 10px;
        font-size: 16px;
        color: black;
        cursor: pointer; /* Makes it look clickable */
    }

    .sidebar-item:hover {
        color: blue; /* Changes color on hover */
    }

    .search-button-container {
        width: 160px;
    }      

    .search-box {
        width: 800px;
    }   
    
""") as demo:

    # Top Bar
    with gr.Row(elem_classes="top-bar"):
        gr.Markdown("**NOMURA**", elem_classes="logo")

        # Profile section in the top-right corner
        gr.HTML("""
            <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.2/css/all.min.css">
            <div class="profile-section">             
                <span class="profile-name">Welcome back!<br>Xavier Lee</span>
                <i class="fas fa-user-circle"></i>  <!-- Profile icon -->
            </div>
        """)

    # Define home and classification pages
    home_page = gr.Column(visible=True)
    classification_page = gr.Column(visible=False)
    search_page = gr.Column(visible=False)
    savedDocuments_page = gr.Column(visible=False)

    # HOME 
    with home_page:
        with gr.Row():
            with gr.Column(scale=1, min_width=200, elem_classes="sidebar"):  # Sidebar column
                with gr.Column():
                    home_btn = gr.Button("🏠 Home")
                    home_btn.click(navigate_home, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    search_btn = gr.Button("🔍 Search")
                    search_btn.click(navigate_search, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    savedDocs_btn = gr.Button("📑 Saved Documents")
                    savedDocs_btn.click(navigate_savedDocuments, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    settings_btn = gr.Button("⚙️ Settings")
                
                # Spacer div to push log out button to the bottom
                with gr.Row(elem_classes="flex-grow"):  
                    pass  

                # Log out button at the bottom of the sidebar
                with gr.Row():
                    gr.Button("Log out")

            with gr.Column(scale=4):
                gr.Markdown("### Upload your Document(s) here")
                file_uploader = gr.File(label="Upload your document", type="filepath")
                chosen_model = gr.Dropdown(label="Choose a model", choices=["gpt3.5-turbo", "deepseek-r1:7b", "llama3.1", "phi4:14b"], value="gpt3.5-turbo", interactive=True)
                notes = gr.Textbox(placeholder="Write any comments you have about your document(s) here.", label="Comments")       
                upload_button = gr.Button("Upload Document(s)")
                output_text = gr.Textbox(label="Upload Status", interactive=False)
                classify_button = gr.Button("Classify Document(s)")
                reset_button = gr.Button("Reset") 

                # resets all entries upon clicking reset button
                reset_button.click(
                    fn=reset_home,
                    outputs=[file_uploader, notes, output_text, chosen_model]
                )

                # Link button to file upload
                upload_button.click(
                    fn=upload_file,
                    inputs=[file_uploader, notes],
                    outputs=[output_text]
                )

                document_preview_output = gr.State()  # Store document preview in state

                classify_button.click(
                    fn=upload_to_classify_preview_document,
                    inputs=[output_text],
                    outputs=[document_preview_output, home_page, classification_page]
                )              

    # CLASSIFICATION 
    ## Need to toggle between diff document previews and classification results for diff documents ##
    with classification_page:
        document_preview_output = gr.State()  # Store document preview in state
        with gr.Row():
            with gr.Column(scale=1, min_width=200, elem_classes="sidebar"):  
                with gr.Column():
                    home_btn = gr.Button("🏠 Home")
                    home_btn.click(navigate_home, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    search_btn = gr.Button("🔍 Search")
                    search_btn.click(navigate_search, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    savedDocs_btn = gr.Button("📑 Saved Documents")
                    savedDocs_btn.click(navigate_savedDocuments, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    settings_btn = gr.Button("⚙️ Settings")
                
                # Spacer div to push log out button to the bottom
                with gr.Row(elem_classes="flex-grow"):  
                    pass  

                # Log out button at the bottom of the sidebar
                with gr.Row():
                    gr.Button("Log out")
            
            with gr.Column(scale=3):
                gr.Markdown("### Document Preview")
                document_preview = gr.Textbox(label="Document Preview", interactive=False)       

                with gr.Row():
                    backtoHome_button = gr.Button("New classification", elem_classes="backtoHome-button") # reset all fields from prev page automatically
                    reclassify_button = gr.Button("Reclassify", elem_classes="reclassify-button")                

            with gr.Column(scale=1):
                gr.Markdown("### Classification Results")
                classification_contentType = gr.TextArea(label="Content type", interactive=False)
                classification_contentType = gr.TextArea(label="Author(s)", interactive=False)
                classification_contentType = gr.TextArea(label="Posted at", interactive=False)

            # Button to return to home page
            backtoHome_button.click(navigate_home, outputs=[home_page, classification_page])
        

    # SEARCH 
    with search_page:
        with gr.Row():
            with gr.Column(scale=1, min_width=200, elem_classes="sidebar"):  # Sidebar column
                with gr.Column():
                    home_btn = gr.Button("🏠 Home")
                    home_btn.click(navigate_home, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    search_btn = gr.Button("🔍 Search")
                    search_btn.click(navigate_search, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    savedDocs_btn = gr.Button("📑 Saved Documents")
                    savedDocs_btn.click(navigate_savedDocuments, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    settings_btn = gr.Button("⚙️ Settings")
                
                # Spacer div to push log out button to the bottom
                with gr.Row(elem_classes="flex-grow"):  
                    pass  

                # Log out button at the bottom of the sidebar
                with gr.Row():
                    gr.Button("Log out")

            with gr.Column(scale=4):
                # Single row containing the search bar and filters
                with gr.Row():
                    search_input = gr.Textbox(label="Search query", placeholder="Enter your search query", elem_classes="search-box", scale=3)
                    filters_contentType = gr.Dropdown(label="Content Type(s)", choices=choices, 
                                                    multiselect=True, elem_classes="filter-box", scale=1)
                    filters_authors = gr.Textbox(label="Author(s)", scale=1, placeholder="Separate names by commas")
                    filters_postedAt = gr.Textbox(label="Posted at", scale=1)
                    search_button = gr.Button("Search", elem_classes="search-button-container", scale=1)

                search_documents_output = gr.TextArea(label="Search Results", interactive=False)

            # Update the search function to consider all filters
            def perform_search(query, content_types, authors, dates):
                return f"Searching for: '{query}'\nFilters Applied:\n- Content Type: {content_types}\n- Authors: {authors}\n- Dates: {dates}"

            # Trigger search with all filters as inputs
            search_button.click(perform_search, 
                inputs=[search_input, filters_contentType, filters_authors, filters_postedAt], 
                outputs=search_documents_output
            )


    # SAVED DOCUMENTS
    with savedDocuments_page:
        with gr.Row():
            with gr.Column(scale=1, min_width=200, elem_classes="sidebar"):  # Sidebar column
                with gr.Column():
                    home_btn = gr.Button("🏠 Home")
                    home_btn.click(navigate_home, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    search_btn = gr.Button("🔍 Search")
                    search_btn.click(navigate_search, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    savedDocs_btn = gr.Button("📑 Saved Documents")
                    savedDocs_btn.click(navigate_savedDocuments, outputs=[home_page, classification_page, search_page, savedDocuments_page])
                    settings_btn = gr.Button("⚙️ Settings")
                
                # Spacer div to push log out button to the bottom
                with gr.Row(elem_classes="flex-grow"):  
                    pass  

                # Log out button at the bottom of the sidebar
                with gr.Row():
                    gr.Button("Log out")

            with gr.Column(scale=4):
                gr.Markdown("### Saved Documents")


   # Link buttons to navigation functions
    home_btn.click(navigate_home, outputs=[home_page, classification_page, search_page, savedDocuments_page])
    classify_button.click(navigate_classification, outputs=[home_page, classification_page, search_page, savedDocuments_page])
    search_btn.click(navigate_search, outputs=[home_page, classification_page, search_page, savedDocuments_page])
    savedDocs_btn.click(navigate_savedDocuments, outputs=[home_page, classification_page, search_page, savedDocuments_page])


if __name__ == "__main__":
    demo.launch()

